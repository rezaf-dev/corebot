#!/usr/bin/env python3
import sys
from docx import Document


def main():
    if len(sys.argv) != 2:
        print("Usage: extract_docx.py /path/to/file.docx", file=sys.stderr)
        return 2

    document = Document(sys.argv[1])
    parts = []

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    text = "\n".join(parts).strip()
    if not text:
        print("No extractable text found", file=sys.stderr)
        return 1

    print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
